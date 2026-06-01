# -*- coding: utf-8 -*-
"""Build FMS Authoritarian Exit Dilemma Academic Article DOCX from markdown."""
from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
INPUT_MD = ROOT / "FMS-Academic-Article-Authoritarian-Exit-Dilemma.md"
OUTPUT_DOCX = ROOT / "FMS-Academic-Article-Authoritarian-Exit-Dilemma.docx"

NAVY = RGBColor(0x00, 0x33, 0x66)
BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)
stats = {"tables": 0, "figures": 0}


def set_run_font(run, name=BODY_FONT, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def setup_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    for level, size in [(1, 14), (2, 12), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = BODY_FONT
        h.font.bold = True
        h.font.color.rgb = NAVY
        h.font.size = Pt(size)
    try:
        doc.styles["Caption"]
    except KeyError:
        cap = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = BODY_FONT
        cap.font.size = Pt(10)
        cap.font.italic = True
    try:
        ref = doc.styles["Bibliography"]
    except KeyError:
        ref = doc.styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    ref.font.name = BODY_FONT
    ref.font.size = Pt(12)
    ref.paragraph_format.left_indent = Inches(0.5)
    ref.paragraph_format.first_line_indent = Inches(-0.5)


def parse_inline(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_run_font(paragraph.add_run(part[2:-2]), bold=True)
        elif part.startswith("*") and part.endswith("*"):
            set_run_font(paragraph.add_run(part[1:-1]), italic=True)
        else:
            set_run_font(paragraph.add_run(part))


def is_table_row(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def parse_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator_row(line):
    return bool(re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", line))


def add_markdown_table(doc, rows, caption_above=None):
    if caption_above:
        cp = doc.add_paragraph(caption_above)
        cp.style = "Caption"
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            parse_inline(cell.paragraphs[0], row[j] if j < len(row) else "")
    stats["tables"] += 1


def add_title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        t.add_run(
            "Authoritarian Exit Dilemmas and Democratic Transition: "
            "The Role of Credible Exit Guarantees in Conflict Prevention"
        ),
        size=Pt(14),
        bold=True,
    )
    doc.add_paragraph()
    for line in (
        "Dr. Fatih Sayin",
        "Foundation for Multilateral Strategies (FMS)",
        "www.fmsthinktank.org",
        "May 2026",
    ):
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(mp.add_run(line), size=Pt(12))
    doc.add_page_break()


def heading_level(line):
    m = re.match(r"^(#{1,6})\s+(.*)$", line)
    return (len(m.group(1)), m.group(2).strip()) if m else None


def build():
    lines = INPUT_MD.read_text(encoding="utf-8", errors="replace").splitlines()
    doc = Document()
    setup_document(doc)
    add_title_page(doc)

    i = 0
    skip_front = True
    in_references = False
    pending_figure_caption = None

    while i < len(lines):
        line = lines[i]

        if skip_front:
            if line.strip().startswith("## Abstract"):
                skip_front = False
            else:
                i += 1
                continue

        hl = heading_level(line)
        if hl:
            level, title = hl
            if title.lower().startswith("references"):
                in_references = True
            if level == 1:
                doc.add_paragraph(title, style="Heading 1")
            elif level == 2:
                doc.add_paragraph(title, style="Heading 2")
            elif level == 3:
                doc.add_paragraph(title, style="Heading 3")
            else:
                p = doc.add_paragraph()
                set_run_font(p.add_run(title), bold=True)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        if is_table_row(line):
            rows = []
            while i < len(lines) and is_table_row(lines[i]):
                if not is_separator_row(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, rows, caption_above=pending_figure_caption)
            pending_figure_caption = None
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if re.match(r"^\*\*Figure", stripped, re.I) or stripped.lower().startswith("figure "):
            pending_figure_caption = stripped.strip("*")
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        if in_references:
            p = doc.add_paragraph(style="Bibliography")
            parse_inline(p, stripped)
        else:
            p = doc.add_paragraph()
            parse_inline(p, stripped)
        i += 1

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    out = build()
    print("OUTPUT", out)
    print("SIZE", out.stat().st_size)
    print("ZIP", zipfile.is_zipfile(out))
    print("TABLES", stats["tables"])

# -*- coding: utf-8 -*-
"""Build FMS Strategic Review flagship DOCX from markdown."""
from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
INPUT_MD = ROOT / "reports" / "FMS-Strategic-Review-AI-Diplomacy-Peace-Publication.md"
OUTPUT_DOCX = ROOT / "reports" / "FMS-Strategic-Review-AI-Diplomacy-Peace-Publication.docx"

NAVY = RGBColor(0x00, 0x33, 0x66)
GOLD = RGBColor(0xC5, 0x9A, 0x2E)
BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)
stats = {"tables": 0, "figures": 0, "mermaid": 0}


def set_run_font(run, name=BODY_FONT, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def setup_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    for level, size in [(1, 16), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = BODY_FONT
        h.font.bold = True
        h.font.color.rgb = NAVY
        h.font.size = Pt(size)
    try:
        doc.styles["Caption"]
    except KeyError:
        cap = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = BODY_FONT
        cap.font.size = Pt(10)
        cap.font.italic = True
    try:
        ref = doc.styles["Bibliography"]
    except KeyError:
        ref = doc.styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    ref.font.name = BODY_FONT
    ref.font.size = Pt(12)
    ref.paragraph_format.left_indent = Inches(0.5)
    ref.paragraph_format.first_line_indent = Inches(-0.5)



def add_field_run(paragraph, instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    run._r.append(instr_el)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    t = OxmlElement("w:t")
    t.text = "1"
    run._r.append(t)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Table of Contents")
    set_run_font(r, size=Pt(16), bold=True, color=NAVY)
    doc.add_paragraph()
    toc_p = doc.add_paragraph()
    add_field_run(toc_p, 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph("(Right-click and Update Field in Word.)")
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)


def add_page_number_footer(section, roman=False, start_at=None):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in list(p.runs):
        p._p.remove(run._r)
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:fmt"), "lowerRoman" if roman else "decimal")
    if start_at is not None:
        pg.set(qn("w:start"), str(start_at))
    add_field_run(p, "PAGE")


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(t.add_run("FMS STRATEGIC REVIEW"), size=Pt(28), bold=True, color=NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub.add_run("Volume 1, Issue 1 (2026)"), size=Pt(14), italic=True, color=GOLD)
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        title.add_run("AI-Powered Diplomacy and the Architecture of Sustainable Peace"),
        size=Pt(18),
        bold=True,
    )
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        st.add_run(
            "How Artificial Intelligence Can Strengthen Preventive Diplomacy, "
            "Multilateral Governance, and Human Dignity"
        ),
        size=Pt(12),
        italic=True,
    )
    for line in (
        "Editor-in-Chief: Dr. Fatih Sayin",
        "Foundation for Multilateral Strategies (FMS)",
        "Publication Date: 29 May 2026",
    ):
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(mp.add_run(line), size=Pt(11))


def parse_inline(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_run_font(paragraph.add_run(part[2:-2]), bold=True)
        elif part.startswith("*") and part.endswith("*"):
            set_run_font(paragraph.add_run(part[1:-1]), italic=True)
        else:
            set_run_font(paragraph.add_run(part))


def is_table_row(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def parse_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator_row(line):
    return bool(re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", line))


def add_markdown_table(doc, rows, caption_above=None):
    if caption_above:
        cp = doc.add_paragraph(caption_above)
        cp.style = "Caption"
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            parse_inline(cell.paragraphs[0], row[j] if j < len(row) else "")
    stats["tables"] += 1


def add_mermaid_figure(doc, caption):
    tbl = doc.add_table(rows=3, cols=3)
    tbl.style = "Table Grid"
    labels = [
        "1. Preventive Diplomacy",
        "2. Human Dignity",
        "3. AI Governance",
        "4. Economic Stability",
        "5. Climate and Resources",
        "6. Education and Youth",
        "7. Multilateral Cooperation",
    ]
    coords = [(0, 1), (1, 0), (1, 1), (1, 2), (2, 0), (2, 1), (2, 2)]
    for (r, c), label in zip(coords, labels):
        tbl.rows[r].cells[c].text = label
    tbl.rows[1].cells[1].text = "GDAIC Hub"
    stats["figures"] += 1
    stats["mermaid"] += 1
    cp = doc.add_paragraph(caption or "Figure 1: FMS Peace Architecture Framework")
    cp.style = "Caption"
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def heading_level(line):
    m = re.match(r"^(#{1,6})\s+(.*)$", line)
    return (len(m.group(1)), m.group(2).strip()) if m else None


def build():
    lines = INPUT_MD.read_text(encoding="utf-8", errors="replace").splitlines()
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    doc.add_page_break()
    add_page_number_footer(doc.sections[0], roman=True, start_at=1)

    i = 0
    skip_until = None
    in_references = False
    pending_table_caption = None
    pending_figure_caption = None
    body_started = False

    while i < len(lines):
        line = lines[i]
        hl = heading_level(line)
        if hl:
            level, title = hl
            upper = title.upper()
            if upper == "COVER PAGE":
                skip_until = "EDITORIAL"
                i += 1
                continue
            if skip_until == "EDITORIAL" and upper.startswith("EDITORIAL"):
                skip_until = None
            if upper.startswith("TABLE OF CONTENTS"):
                doc.add_page_break()
                add_toc(doc)
                doc.add_page_break()
                if not body_started:
                    doc.add_section(WD_SECTION.NEW_PAGE)
                    add_page_number_footer(doc.sections[-1], roman=False, start_at=1)
                    body_started = True
                skip_until = "EXEC"
                i += 1
                continue
            if skip_until == "EXEC" and title.upper().startswith("1. EXECUTIVE"):
                skip_until = None
            if skip_until:
                i += 1
                continue
            if level == 1:
                if re.match(r"^21\.|^REFERENCES", title, re.I):
                    in_references = True
                elif re.match(r"^22\.|^APPENDICES", title, re.I):
                    in_references = False
                if body_started:
                    doc.add_page_break()
                doc.add_paragraph(title, style=f"Heading {min(level, 3)}")
                i += 1
                continue

        if skip_until:
            i += 1
            continue
        if line.strip() == "---":
            i += 1
            continue
        if line.strip().startswith("```"):
            if line.strip().startswith("```mermaid"):
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    i += 1
                i += 1
                add_mermaid_figure(doc, pending_figure_caption)
                pending_figure_caption = None
                continue
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            set_run_font(p.add_run("\n".join(code)), name="Courier New", size=Pt(10))
            continue
        if is_table_row(line):
            rows = []
            while i < len(lines) and is_table_row(lines[i]):
                if not is_separator_row(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, rows, caption_above=pending_table_caption)
            pending_table_caption = None
            continue
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if re.match(r"^\*\*Figure", stripped, re.I) or stripped.lower().startswith("figure "):
            pending_figure_caption = stripped.strip("*")
            i += 1
            continue
        if re.match(r"^\*\*Table", stripped, re.I) or stripped.lower().startswith("table "):
            pending_table_caption = stripped.strip("*")
            i += 1
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, stripped[2:].strip())
            i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue
        if in_references:
            p = doc.add_paragraph(style="Bibliography")
            parse_inline(p, stripped)
        else:
            p = doc.add_paragraph()
            parse_inline(p, stripped)
        i += 1

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    out = build()
    print("OUTPUT", out)
    print("SIZE", out.stat().st_size)
    print("ZIP", zipfile.is_zipfile(out))
    print("TABLES", stats["tables"])
    print("FIGURES", stats["figures"])
    print("MERMAID", stats["mermaid"])